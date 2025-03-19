from typing import Any
import httpx
import os
import re
import subprocess
import tempfile
import uuid
from pathlib import Path
import requests
import json
import base64
from io import BytesIO
from bs4 import BeautifulSoup

import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


def render_mermaid_to_image(mermaid_code, output_path=None):
    """
    Render Mermaid diagram to an image using multiple methods.
    Returns the path to the saved image.
    """
    # Create a temporary file to save the image if not provided
    if not output_path:
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, f"mermaid_{uuid.uuid4()}.png")
    
    # Method 1: Try using Mermaid.ink API
    try:
        # Encode the Mermaid code for the URL
        encoded_data = {"code": mermaid_code}
        json_str = json.dumps(encoded_data)
        base64_str = base64.urlsafe_b64encode(json_str.encode('utf-8')).decode('utf-8')
        
        # Use the Mermaid.ink API
        api_url = f"https://mermaid.ink/img/{base64_str}"
        response = requests.get(api_url)
        
        if response.status_code == 200:
            # Save the image
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path
    except Exception as e:
        print(f"Error using Mermaid.ink API: {e}")
    
    # Method 2: Try using mermaid-cli if available
    try:
        # Check if mmdc (mermaid-cli) is installed
        subprocess.run(["mmdc", "--version"], capture_output=True, check=True)
        
        # Create a temporary file for the Mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as temp_mmd:
            temp_mmd.write(mermaid_code)
            temp_mmd_path = temp_mmd.name
        
        # Run mmdc to generate the image
        subprocess.run([
            "mmdc",
            "-i", temp_mmd_path,
            "-o", output_path,
            "-b", "transparent"
        ], check=True)
        
        # Clean up the temporary Mermaid file
        os.unlink(temp_mmd_path)
        
        return output_path
    except (subprocess.SubprocessError, FileNotFoundError) as e:
        print(f"Error using mermaid-cli: {e}")
    
    # Method 3: Try using the Kroki API
    try:
        payload = {
            "diagram_source": mermaid_code,
            "diagram_type": "mermaid",
            "output_format": "png"
        }
        
        response = requests.post("https://kroki.io/", json=payload)
        
        if response.status_code == 200:
            # Save the image
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path
    except Exception as e:
        print(f"Error using Kroki API: {e}")
    
    # All methods failed
    print("All rendering methods failed for Mermaid diagram")
    return None


def extract_mermaid_blocks(md_content):
    """Extract Mermaid code blocks from Markdown content."""
    # Pattern to match ```mermaid ... ``` blocks
    pattern = r'```mermaid\s+(.*?)\s+```'
    # Find all matches using re.DOTALL to match across multiple lines
    matches = re.findall(pattern, md_content, re.DOTALL)
    return matches


def extract_code_blocks(md_content):
    """Extract all code blocks (except Mermaid) from Markdown content."""
    # 使用更精确的模式匹配代码块
    # 匹配 ```language 和 ``` 之间的内容
    pattern = r'```(?!mermaid)([^\n]*)\n([\s\S]*?)\n```'
    matches = re.findall(pattern, md_content, re.DOTALL)
    
    # 处理匹配结果，清理语言和代码内容
    cleaned_matches = []
    for lang, code in matches:
        # 清理语言标识
        lang = lang.strip()
        # 清理代码内容，但保留换行符
        code = code.rstrip()  # 只去除尾部空白，保留换行
        cleaned_matches.append((lang, code))
    
    print(f"Found {len(cleaned_matches)} code blocks")
    for i, (lang, code) in enumerate(cleaned_matches):
        print(f"Code block {i}: language='{lang}', length={len(code)} chars")
        if len(code) > 50:
            preview = code.split('\n')[0][:50]  # 只显示第一行的前50个字符
            print(f"  Preview (first line): {preview}...")
        else:
            print(f"  Content: {code}")
    
    return cleaned_matches


def html_to_docx(html_content, doc, table_style='Table Grid'):
    """Convert HTML content to Word document elements."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 首先检查是否有代码块元素
    code_elements = soup.find_all('code')
    if code_elements:
        print(f"Found {len(code_elements)} code elements in HTML")
    
    pre_elements = soup.find_all('pre')
    if pre_elements:
        print(f"Found {len(pre_elements)} pre elements in HTML")
        # 处理pre元素（通常包含代码块）
        for pre in pre_elements:
            code = pre.get_text(strip=True)
            if code:
                print(f"Processing pre element with content length: {len(code)}")
                # 检查是否有语言类
                lang = ""
                if pre.has_attr('class'):
                    for cls in pre['class']:
                        if cls.startswith('language-'):
                            lang = cls.replace('language-', '')
                            break
                
                # 使用代码块格式化函数
                format_code_block(doc, code, lang)
                # 从soup中移除已处理的元素，避免重复处理
                pre.extract()
    
    # Process elements in order
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote', 'table']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            paragraph.style = f'Heading {element.name[1]}'
        
        elif element.name == 'p':
            # 检查段落中是否包含代码元素
            code_in_p = element.find('code')
            if code_in_p and len(element.contents) == 1:
                # 如果段落只包含一个代码元素，将其作为代码块处理
                code = code_in_p.get_text(strip=True)
                lang = ""
                if code_in_p.has_attr('class'):
                    for cls in code_in_p['class']:
                        if cls.startswith('language-'):
                            lang = cls.replace('language-', '')
                            break
                format_code_block(doc, code, lang)
            else:
                # 正常处理段落，不添加边框
                paragraph = doc.add_paragraph(element.get_text(strip=True))
                apply_style_to_paragraph(paragraph, element)
        
        elif element.name == 'blockquote':
            paragraph = doc.add_paragraph(element.get_text(strip=True))
            paragraph.style = 'Quote'
        
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                process_list_item(doc, li, 'List Bullet')
        
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                process_list_item(doc, li, 'List Number')
        
        elif element.name == 'table':
            process_table(doc, element, table_style)
    
    return doc


def apply_style_to_paragraph(paragraph, element):
    """Apply HTML styles to a Word paragraph based on the element."""
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(element.name[1])
        paragraph.style = f'Heading {level}'
    
    if element.name == 'strong' or element.find('strong'):
        for run in paragraph.runs:
            run.bold = True
    
    if element.name == 'em' or element.find('em'):
        for run in paragraph.runs:
            run.italic = True
    
    if element.name == 'u' or element.find('u'):
        for run in paragraph.runs:
            run.underline = True
    
    if element.name == 'code' or element.find('code'):
        for run in paragraph.runs:
            run.font.name = 'Courier New'
    
    if element.name == 'center' or element.get('align') == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def process_list_item(doc, li_element, list_style, level=0):
    """Process a list item and its children with proper indentation."""
    # Add the list item with proper style and level
    text = li_element.get_text(strip=True)
    paragraph = doc.add_paragraph(text)
    paragraph.style = list_style
    paragraph.paragraph_format.left_indent = Pt(18 * level)  # Indent based on nesting level
    
    # Process any nested lists
    nested_ul = li_element.find('ul')
    nested_ol = li_element.find('ol')
    
    if nested_ul:
        for nested_li in nested_ul.find_all('li', recursive=False):
            process_list_item(doc, nested_li, 'List Bullet', level + 1)
    
    if nested_ol:
        for nested_li in nested_ol.find_all('li', recursive=False):
            process_list_item(doc, nested_li, 'List Number', level + 1)


def process_table(doc, table_element, table_style='Table Grid'):
    """Process a table element and convert it to a Word table."""
    # Find all rows in the table
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Count the maximum number of cells in any row
    max_cols = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        max_cols = max(max_cols, len(cells))
    
    if max_cols == 0:
        return
    
    # Create the table in the document
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = table_style
    
    # Fill the table with data
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < max_cols:  # Ensure we don't exceed the table dimensions
                # Get cell text and apply basic formatting
                text = cell.get_text(strip=True)
                table.cell(i, j).text = text
                
                # Apply header formatting if it's a header cell
                if cell.name == 'th' or i == 0:
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def add_border_to_paragraph(paragraph):
    """Add border to a paragraph."""
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    # The order of these elements matters for Word
    for border_pos in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_pos}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width in 1/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        pBdr.append(border)
    
    # Insert the border element before these elements
    child_elements = [
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    ]
    
    # Find the first child in the list that exists
    insert_after = None
    for child in child_elements:
        element = pPr.find(qn(child))
        if element is not None:
            insert_after = element
            break
    
    # If none of the specified children is found, just append to the end
    if insert_after is None:
        pPr.append(pBdr)
    else:
        insert_after.addprevious(pBdr)


def format_code_block(doc, code, language=""):
    """Format a code block with proper styling in the Word document."""
    print(f"Formatting code block with language: '{language}', code length: {len(code)} chars")
    
    # 创建一个段落用于代码块
    code_para = doc.add_paragraph()
    code_para.style = 'No Spacing'
    
    # 如果有语言标识，添加到代码块前面
    if language.strip():
        lang_run = code_para.add_run(f"Language: {language}\n")
        lang_run.bold = True
        lang_run.font.size = Pt(9)
    
    # 添加代码文本，保留换行符
    # 将代码按行分割，然后逐行添加，保留换行
    lines = code.split('\n')
    for i, line in enumerate(lines):
        if i > 0:  # 不是第一行，先添加换行符
            code_para.add_run('\n')
        code_run = code_para.add_run(line)
        code_run.font.name = 'Times New Roman'
        code_run.font.size = Pt(10)
    
    # 添加边框
    try:
        add_border_to_paragraph(code_para)
        print("Successfully added border to code paragraph")
    except Exception as e:
        print(f"Error adding border to code paragraph: {e}")
        # 如果添加边框失败，尝试使用另一种方法
        try:
            p = code_para._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            
            for border_pos in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_pos}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                pBdr.append(border)
            
            pPr.append(pBdr)
            print("Successfully added border using alternative method")
        except Exception as e2:
            print(f"Alternative border method also failed: {e2}")
    
    # 添加一些空间在代码块后面
    doc.add_paragraph()
    
    return doc


def check_and_process_code_blocks(md_content, doc):
    """直接从Markdown内容中提取和处理代码块，不依赖于HTML转换。"""
    print("直接从Markdown内容中提取和处理代码块...")
    
    # 使用精确的正则表达式匹配代码块
    code_blocks = re.findall(r'```([^\n]*)\n([\s\S]*?)\n```', md_content, re.DOTALL)
    
    if not code_blocks:
        print("未找到代码块")
        return False
    
    print(f"找到 {len(code_blocks)} 个代码块")
    
    # 处理每个代码块
    for i, (lang, code) in enumerate(code_blocks):
        lang = lang.strip()
        code = code.strip()
        
        print(f"处理代码块 {i+1}: 语言='{lang}', 长度={len(code)} 字符")
        
        # 跳过mermaid代码块，因为它们会被单独处理
        if "mermaid" not in lang.lower():
            doc.add_paragraph(f"代码块 {i+1}:", style='Heading 3')
            format_code_block(doc, code, lang)
    
    return True


def debug_code_blocks(md_content):
    """详细分析Markdown内容中的代码块，用于调试"""
    print("\n===== 代码块调试信息 =====")
    
    # 1. 使用简单模式查找所有代码块（包括mermaid）
    simple_pattern = r'```(.*?)```'
    simple_blocks = re.findall(simple_pattern, md_content, re.DOTALL)
    print(f"简单模式找到 {len(simple_blocks)} 个代码块")
    
    # 2. 使用精确模式查找所有代码块
    precise_pattern = r'```([^\n]*)\n([\s\S]*?)\n```'
    precise_blocks = re.findall(precise_pattern, md_content, re.DOTALL)
    print(f"精确模式找到 {len(precise_blocks)} 个代码块")
    
    # 3. 分析每个代码块
    for i, (lang, code) in enumerate(precise_blocks):
        lang = lang.strip()
        # 检查代码块是否包含特殊字符
        special_chars = [c for c in code if ord(c) > 127]
        if special_chars:
            print(f"  包含 {len(special_chars)} 个特殊字符")
            print(f"  特殊字符: {special_chars[:10]}")
        
        # 检查代码块是否包含可能导致问题的字符
        problem_chars = ['\\', '{', '}', '[', ']', '(', ')', '$', '^', '*', '+', '?', '.', '|']
        found_problem_chars = [c for c in problem_chars if c in code]
        if found_problem_chars:
            print(f"  包含可能导致正则表达式问题的字符: {found_problem_chars}")
    
    # 4. 尝试提取mermaid代码块
    mermaid_pattern = r'```mermaid\s+(.*?)\s+```'
    mermaid_blocks = re.findall(mermaid_pattern, md_content, re.DOTALL)
    print(f"\n找到 {len(mermaid_blocks)} 个mermaid代码块")
    
    # 5. 尝试提取非mermaid代码块
    non_mermaid_pattern = r'```(?!mermaid)([^\n]*)\n([\s\S]*?)\n```'
    non_mermaid_blocks = re.findall(non_mermaid_pattern, md_content, re.DOTALL)
    print(f"找到 {len(non_mermaid_blocks)} 个非mermaid代码块")
    
    print("===== 代码块调试信息结束 =====\n")
    return precise_blocks


def debug_code_block_processing(md_content, output_file="debug_code_blocks.docx", table_style='Table Grid'):
    """专门用于调试代码块处理的函数"""
    print("\n===== 开始调试代码块处理 =====")
    
    # 创建一个新的Word文档
    doc = Document()
    doc.add_paragraph("代码块处理调试报告", style='Title')
    
    # 1. 添加原始Markdown内容
    doc.add_paragraph("原始Markdown内容", style='Heading 1')
    p = doc.add_paragraph(md_content)
    p.style = 'No Spacing'
    
    # 2. 提取并分析代码块
    doc.add_paragraph("代码块分析", style='Heading 1')
    
    # 使用不同的正则表达式模式提取代码块
    patterns = [
        ("简单模式", r'```(.*?)```'),
        ("精确模式", r'```([^\n]*)\n([\s\S]*?)\n```'),
        ("Mermaid模式", r'```mermaid\s+(.*?)\s+```'),
        ("非Mermaid模式", r'```(?!mermaid)([^\n]*)\n([\s\S]*?)\n```')
    ]
    
    for name, pattern in patterns:
        doc.add_paragraph(f"{name} 提取结果", style='Heading 2')
        
        try:
            if name in ["简单模式", "Mermaid模式"]:
                matches = re.findall(pattern, md_content, re.DOTALL)
                doc.add_paragraph(f"找到 {len(matches)} 个匹配")
                
                for i, match in enumerate(matches):
                    doc.add_paragraph(f"匹配 {i+1}:", style='Heading 3')
                    code_para = doc.add_paragraph(match)
                    code_para.style = 'No Spacing'
                    add_border_to_paragraph(code_para)
            else:
                matches = re.findall(pattern, md_content, re.DOTALL)
                doc.add_paragraph(f"找到 {len(matches)} 个匹配")
                
                for i, (lang, code) in enumerate(matches):
                    doc.add_paragraph(f"匹配 {i+1}:", style='Heading 3')
                    doc.add_paragraph(f"语言: {lang}")
                    code_para = doc.add_paragraph(code)
                    code_para.style = 'No Spacing'
                    add_border_to_paragraph(code_para)
        except Exception as e:
            doc.add_paragraph(f"提取过程出错: {str(e)}")
    
    # 3. 测试代码块格式化
    doc.add_paragraph("代码块格式化测试", style='Heading 1')
    
    # 提取所有代码块
    try:
        precise_blocks = re.findall(r'```([^\n]*)\n([\s\S]*?)\n```', md_content, re.DOTALL)
        
        for i, (lang, code) in enumerate(precise_blocks):
            doc.add_paragraph(f"代码块 {i+1} 格式化:", style='Heading 2')
            format_code_block(doc, code, lang.strip())
    except Exception as e:
        doc.add_paragraph(f"格式化过程出错: {str(e)}")
    
    # 保存文档
    doc.save(output_file)
    print(f"调试报告已保存到: {output_file}")
    print("===== 代码块处理调试结束 =====\n")
    
    return output_file


def preprocess_markdown(md_content):
    """预处理Markdown内容，确保表格前后有空行"""
    # 使用正则表达式查找表格并确保其前后有空行
    # 表格通常以 | 开头的行为标志
    lines = md_content.split('\n')
    processed_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        processed_lines.append(line)
        
        # 检测表格开始
        if line.strip().startswith('|') and i + 1 < len(lines) and '|' in lines[i+1] and '-' in lines[i+1]:
            # 如果前一行不是空行，添加一个空行
            if i > 0 and processed_lines[-2].strip() != '':
                processed_lines.insert(-1, '')
            
            # 添加表格行直到表格结束
            table_end_idx = i
            while table_end_idx + 1 < len(lines) and lines[table_end_idx + 1].strip().startswith('|'):
                table_end_idx += 1
                processed_lines.append(lines[table_end_idx])
                
            # 如果表格后没有空行，添加一个空行
            if table_end_idx + 1 < len(lines) and lines[table_end_idx + 1].strip() != '':
                processed_lines.append('')
                
            i = table_end_idx
        
        i += 1
    
    return '\n'.join(processed_lines)



def md_to_docx(md_content, output_file=None, debug_mode=False, table_style='Table Grid'):
    """Convert Markdown content to a DOCX file, rendering Mermaid diagrams as images.
    
    Args:
        md_content: Markdown content to convert
        output_file: Optional output file path, defaults to 'output.docx'
        debug_mode: Whether to enable debug mode
        table_style: Style to apply to tables, defaults to 'Table Grid'
                     Common styles include: 'Table Normal', 'Table Grid', 'Light Shading',
                     'Light List', 'Light Grid', 'Medium Shading 1', 'Medium Shading 2',
                     'Medium List 1', 'Medium List 2', 'Medium Grid 1', 'Medium Grid 2',
                     'Medium Grid 3', 'Dark List', 'Colorful List', 'Colorful Grid',
                     'Light Shading Accent 1', 'Light List Accent 1', etc.
        
    Returns:
        The path to the saved DOCX file
    """
    if debug_mode:
        print("\n===== 开始调试模式 =====")
        debug_code_blocks(md_content)
    
    print(f"Processing Markdown content of length: {len(md_content)} chars")
    
    # 检查是否有代码块 - 使用更精确的模式
    code_block_pattern = r'```([^\n]*)\n([\s\S]*?)\n```'
    all_code_blocks = re.findall(code_block_pattern, md_content, re.DOTALL)
    code_blocks = []
    print(f"Total code blocks found (simple check): {len(all_code_blocks)}")
    for i, (lang, code) in enumerate(all_code_blocks):
        print(f"Simple check - Block {i}: language='{lang.strip()}', length={len(code)} chars")
        if lang.strip() != "mermaid":
            code_blocks.append((lang.strip(), code.strip()))
    # Extract Mermaid blocks
    mermaid_blocks = extract_mermaid_blocks(md_content)
    print(f"Found {len(mermaid_blocks)} Mermaid blocks")
    
    # Create a new Word document
    doc = Document()
    
    md_content = preprocess_markdown(md_content)

    # 创建一个副本用于替换操作
    md_content_copy = md_content
    
    # Replace Mermaid blocks with placeholders and keep track of them
    mermaid_placeholders = []
    for i, block in enumerate(mermaid_blocks):
        placeholder = f"MERMAID_DIAGRAM_{i}"
        mermaid_placeholders.append(placeholder)
        # Use regex to replace the block to handle different whitespace patterns
        pattern = r'```mermaid\s+' + re.escape(block) + r'\s+```'
        md_content_copy = re.sub(pattern, placeholder, md_content_copy, flags=re.DOTALL)
    
    # Replace code blocks with placeholders and keep track of them
    code_placeholders = []
    for i, (language, block) in enumerate(code_blocks):
        placeholder = f"CODE_BLOCK_{i}"
        code_placeholders.append((placeholder, language, block))
        
        # 使用更精确的模式来替换代码块，保留换行
        if language:
            # 有语言标识的代码块
            pattern = r'```' + re.escape(language) + r'\n' + re.escape(block) + r'\n```'
        else:
            # 无语言标识的代码块
            pattern = r'```\n' + re.escape(block) + r'\n```'
        
        # 尝试替换
        before_replace = md_content_copy
        md_content_copy = re.sub(pattern, placeholder, md_content_copy, count=1, flags=re.DOTALL)
        
        # 如果第一次替换失败，尝试使用更宽松的模式
        if md_content_copy == before_replace:
            print(f"First replacement attempt failed for block {i}, trying looser pattern...")
            if language:
                pattern = r'```' + re.escape(language) + r'[\s\S]*?' + re.escape(block) + r'[\s\S]*?```'
            else:
                pattern = r'```[\s\S]*?' + re.escape(block) + r'[\s\S]*?```'
            
            md_content_copy = re.sub(pattern, placeholder, md_content_copy, count=1, flags=re.DOTALL)
        
        # 检查是否成功替换
        if md_content_copy == before_replace:
            print(f"Warning: Failed to replace code block {i} with language '{language}'")
            # 如果替换失败，尝试直接使用代码内容作为标识符
            md_content_copy = md_content_copy.replace(block, placeholder)
            if md_content_copy == before_replace:
                print(f"Warning: Also failed with direct content replacement for block {i}")
            else:
                print(f"Success: Replaced code block {i} using direct content")
        else:
            print(f"Success: Replaced code block {i} using pattern")
    
    # Combine all placeholders
    all_placeholders = mermaid_placeholders + [p[0] for p in code_placeholders]
    print(f"Total placeholders: {len(all_placeholders)}")
    
    # 如果没有找到任何代码块，但原始内容中确实有代码块标记，则尝试直接处理
    if len(code_blocks) == 0 and '```' in md_content:
        print("No code blocks extracted but found code block markers. Trying direct processing...")
        # 直接在文档中添加一个示例代码块
        doc.add_paragraph("以下是从Markdown中提取的代码块：", style='Heading 2')
        
        # 尝试直接匹配和处理代码块
        direct_code_blocks = re.findall(r'```([^\n]*)\n([\s\S]*?)\n```', md_content, re.DOTALL)
        for i, (lang, code) in enumerate(direct_code_blocks):
            lang = lang.strip()
            # 保留代码中的换行符，不要去除
            code = code.rstrip()  # 只去除尾部空白，保留换行
            print(f"Direct processing - Code block {i}: language='{lang}', length={len(code)} chars")
            if "mermaid" not in lang.lower():  # 跳过mermaid代码块
                format_code_block(doc, code, lang)
    
    # Convert Markdown to HTML with extensions
    html_content = markdown.markdown(
        md_content_copy,
        extensions=[
            'markdown.extensions.extra',
            'markdown.extensions.codehilite',
            'markdown.extensions.tables',
            'markdown.extensions.toc'
        ]
    )
    
    # Split HTML by placeholders
    parts = []
    if all_placeholders:
        print("Splitting HTML content by placeholders...")
        for part in re.split(f"({'|'.join(all_placeholders)})", html_content):
            if part in mermaid_placeholders:
                # This is a mermaid placeholder
                idx = mermaid_placeholders.index(part)
                parts.append(('mermaid', idx))
                print(f"Found mermaid placeholder: {part} (index {idx})")
            elif any(part == p[0] for p in code_placeholders):
                # This is a code block placeholder
                for i, (placeholder, lang, code) in enumerate(code_placeholders):
                    if part == placeholder:
                        parts.append(('code', i))
                        print(f"Found code placeholder: {part} (index {i})")
                        break
            else:
                # This is regular HTML content
                parts.append(('html', part))
                if part.strip():
                    preview = part[:50] + "..." if len(part) > 50 else part
                    print(f"Found HTML content: {preview}")
    else:
        # 如果没有占位符，则将整个内容作为HTML处理
        parts.append(('html', html_content))
        print("No placeholders found, treating entire content as HTML")
    
    print(f"Total parts after splitting: {len(parts)}")
    
    # Process each part
    part_count = {'mermaid': 0, 'code': 0, 'html': 0}
    for part_type, content in parts:
        part_count[part_type] += 1
        
        if part_type == 'mermaid':
            # Render the Mermaid diagram
            mermaid_code = mermaid_blocks[content]
            img_path = render_mermaid_to_image(mermaid_code)
            
            if img_path:
                # Add the image to the document
                doc.add_picture(img_path, width=Inches(6))
                
                # Clean up the temporary image file
                try:
                    os.unlink(img_path)
                except:
                    pass
            else:
                # If rendering failed, add the Mermaid code as text
                doc.add_paragraph("Failed to render Mermaid diagram:", style='Intense Quote')
                code_para = doc.add_paragraph(mermaid_code)
                code_para.style = 'No Spacing'
                for run in code_para.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        
        elif part_type == 'code':
            # Format the code block
            print(f"Processing code block {content}")
            language, code = code_placeholders[content][1], code_placeholders[content][2]
            print(f"Code block language: '{language}', length: {len(code)} chars")
            # 确保代码中的换行被保留
            format_code_block(doc, code, language)
            
        elif part_type == 'html':
            # Add regular content as paragraphs with proper formatting
            if content.strip():
                html_to_docx(content, doc, table_style)
    
    print(f"Processed parts: {part_count}")
    
    # Determine output file name if not provided
    if not output_file:
        output_file = 'output.docx'
    
    # Save the document
    doc.save(output_file)
    
    # 如果没有处理任何代码块，但原始内容中有代码块标记，则创建一个新文档直接处理
    if part_count['code'] == 0 and '```' in md_content:
        print("No code blocks were processed in the main flow. Creating a backup document with direct processing.")
        backup_file = output_file.replace('.docx', '_with_code.docx') if output_file.endswith('.docx') else f"{output_file}_with_code.docx"
        
        # 创建一个新文档
        backup_doc = Document()
        backup_doc.add_paragraph("代码块直接处理版本", style='Title')
        
        # 直接处理代码块
        check_and_process_code_blocks(md_content, backup_doc)
        
        # 保存备份文档
        backup_doc.save(backup_file)
        print(f"Backup document saved as: {backup_file}")
        
        # 在原始文档中也尝试直接处理代码块
        doc.add_paragraph("代码块", style='Heading 1')
        check_and_process_code_blocks(md_content, doc)
        
        # 重新保存原始文档
        doc.save(output_file)
        print(f"Updated original document with direct code block processing")
    
    return output_file 