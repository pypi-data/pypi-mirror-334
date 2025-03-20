import os
import numpy as np
import faiss
from typing import List, Tuple, Optional, Dict
import pickle
from dataclasses import dataclass
from tqdm import tqdm
import fitz  # PyMuPDF for PDF files
from docx import Document as DocxDocument  # python-docx for DOCX files
from pathlib import Path

from yaspin import yaspin
from jarvis.jarvis_platform.registry import PlatformRegistry
import lzma  # 添加 lzma 导入
from threading import Lock
import hashlib

from jarvis.jarvis_utils.config import get_max_paragraph_length, get_max_token_count, get_min_paragraph_length, get_thread_count
from jarvis.jarvis_utils.embedding import get_context_token_count, get_embedding, get_embedding_batch, load_embedding_model
from jarvis.jarvis_utils.output import OutputType, PrettyOutput
from jarvis.jarvis_utils.utils import  get_file_md5, init_env, init_gpu_config

@dataclass
class Document:
    """Document class, for storing document content and metadata"""
    content: str  # Document content
    metadata: Dict  # Metadata (file path, position, etc.)
    md5: str = ""  # File MD5 value, for incremental update detection

class FileProcessor:
    """Base class for file processor"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        """Determine if the file can be processed"""
        raise NotImplementedError
        
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract file text content"""
        raise NotImplementedError

class TextFileProcessor(FileProcessor):
    """Text file processor"""
    ENCODINGS = ['utf-8', 'gbk', 'gb2312', 'latin1']
    SAMPLE_SIZE = 8192  # Read the first 8KB to detect encoding
    
    @staticmethod
    def can_handle(file_path: str) -> bool:
        """Determine if the file is a text file by trying to decode it"""
        try:
            # Read the first part of the file to detect encoding
            with open(file_path, 'rb') as f:
                sample = f.read(TextFileProcessor.SAMPLE_SIZE)
                
            # Check if it contains null bytes (usually represents a binary file)
            if b'\x00' in sample:
                return False
                
            # Check if it contains too many non-printable characters (usually represents a binary file)
            non_printable = sum(1 for byte in sample if byte < 32 and byte not in (9, 10, 13))  # tab, newline, carriage return
            if non_printable / len(sample) > 0.3:  # If non-printable characters exceed 30%, it is considered a binary file
                return False
                
            # Try to decode with different encodings
            for encoding in TextFileProcessor.ENCODINGS:
                try:
                    sample.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue
                    
            return False
            
        except Exception:
            return False
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text content, using the detected correct encoding"""
        detected_encoding = None
        try:
            # First try to detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                
            # Try different encodings
            for encoding in TextFileProcessor.ENCODINGS:
                try:
                    raw_data.decode(encoding)
                    detected_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
                    
            if not detected_encoding:
                raise UnicodeDecodeError(f"Failed to decode file with supported encodings: {file_path}") # type: ignore
                
            # Use the detected encoding to read the file
            with open(file_path, 'r', encoding=detected_encoding, errors='replace') as f:
                content = f.read()
                
            # Normalize Unicode characters
            import unicodedata
            content = unicodedata.normalize('NFKC', content)
            
            return content
            
        except Exception as e:
            raise Exception(f"Failed to read file: {str(e)}")

class PDFProcessor(FileProcessor):
    """PDF file processor"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.pdf'
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        text_parts = []
        with fitz.open(file_path) as doc: # type: ignore
            for page in doc:
                text_parts.append(page.get_text()) # type: ignore
        return "\n".join(text_parts)

class DocxProcessor(FileProcessor):
    """DOCX file processor"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.docx'
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

class RAGTool:
    def __init__(self, root_dir: str):
        """Initialize RAG tool
        
        Args:
            root_dir: Project root directory
        """
        with yaspin(text="初始化环境...", color="cyan") as spinner:
            init_env()
            self.root_dir = root_dir
            os.chdir(self.root_dir)
            spinner.text = "环境初始化完成"
            spinner.ok("✅")
        
        # Initialize configuration
        with yaspin(text="初始化配置...", color="cyan") as spinner:
            self.min_paragraph_length = get_min_paragraph_length()  # Minimum paragraph length
            self.max_paragraph_length = get_max_paragraph_length()  # Maximum paragraph length
            self.context_window = 5  # Fixed context window size
            self.max_token_count = int(get_max_token_count() * 0.8)
            spinner.text = "配置初始化完成"
            spinner.ok("✅")
        
        # Initialize data directory
        with yaspin(text="初始化数据目录...", color="cyan") as spinner:
            self.data_dir = os.path.join(self.root_dir, ".jarvis/rag")
            if not os.path.exists(self.data_dir):
                os.makedirs(self.data_dir)
            spinner.text = "数据目录初始化完成"
            spinner.ok("✅")
            
        # Initialize embedding model
        with yaspin(text="初始化模型...", color="cyan") as spinner:
            try:
                self.embedding_model = load_embedding_model()
                self.vector_dim = self.embedding_model.get_sentence_embedding_dimension()
                spinner.text = "模型加载完成"
                spinner.ok("✅")
            except Exception as e:
                spinner.text = "模型加载失败"
                spinner.fail("❌")
                raise

        with yaspin(text="初始化缓存目录...", color="cyan") as spinner:
            self.cache_dir = os.path.join(self.data_dir, "cache")
            if not os.path.exists(self.cache_dir):
                os.makedirs(self.cache_dir)
                
            self.documents: List[Document] = []
            self.index = None
            self.flat_index = None
            self.file_md5_cache = {}
            spinner.text = "缓存目录初始化完成"
            spinner.ok("✅")
        
        # 加载缓存索引
        self._load_cache_index()

        # Register file processors
        with yaspin(text="初始化文件处理器...", color="cyan") as spinner:
            self.file_processors = [
                TextFileProcessor(),
                PDFProcessor(),
                DocxProcessor()
            ]
            spinner.text = "文件处理器初始化完成"
            spinner.ok("✅")


        # Add thread related configuration
        with yaspin(text="初始化线程配置...", color="cyan") as spinner:
            self.thread_count = get_thread_count()
            self.vector_lock = Lock()  # Protect vector list concurrency
            spinner.text = "线程配置初始化完成"
            spinner.ok("✅")

        # 初始化 GPU 内存配置
        with yaspin(text="初始化 GPU 内存配置...", color="cyan") as spinner:
            self.gpu_config = init_gpu_config()
            spinner.text = "GPU 内存配置初始化完成"
            spinner.ok("✅")


    def _get_cache_path(self, file_path: str) -> str:
        """Get cache file path for a document
        
        Args:
            file_path: Original file path
            
        Returns:
            str: Cache file path
        """
        # 使用文件路径的哈希作为缓存文件名
        file_hash = hashlib.md5(file_path.encode()).hexdigest()
        return os.path.join(self.cache_dir, f"{file_hash}.cache")

    def _load_cache_index(self):
        """Load cache index"""
        index_path = os.path.join(self.data_dir, "index.pkl")
        if os.path.exists(index_path):
            try:
                with yaspin(text="加载缓存索引...", color="cyan") as spinner:
                    with lzma.open(index_path, 'rb') as f:
                        cache_data = pickle.load(f)
                        self.file_md5_cache = cache_data.get("file_md5_cache", {})
                    spinner.text = "缓存索引加载完成"
                    spinner.ok("✅")
                        
                # 从各个缓存文件加载文档
                with yaspin(text="加载缓存文件...", color="cyan") as spinner:
                    for file_path in self.file_md5_cache:
                        cache_path = self._get_cache_path(file_path)
                        if os.path.exists(cache_path):
                            try:
                                with lzma.open(cache_path, 'rb') as f:
                                    file_cache = pickle.load(f)
                                    self.documents.extend(file_cache["documents"])
                                spinner.write(f"✅ 加载缓存文件: {file_path}")
                            except Exception as e:
                                spinner.write(f"❌ 加载缓存文件失败: {file_path}: {str(e)}")
                    spinner.text = "缓存文件加载完成"
                    spinner.ok("✅")
                
                # 重建向量索引

                if self.documents:
                    with yaspin(text="重建向量索引...", color="cyan") as spinner:
                        vectors = []
                        for doc in self.documents:
                            cache_path = self._get_cache_path(doc.metadata['file_path'])
                            if os.path.exists(cache_path):
                                with lzma.open(cache_path, 'rb') as f:
                                    file_cache = pickle.load(f)
                                    doc_idx = next((i for i, d in enumerate(file_cache["documents"]) 
                                                if d.metadata['chunk_index'] == doc.metadata['chunk_index']), None)
                                    if doc_idx is not None:
                                        vectors.append(file_cache["vectors"][doc_idx])
                        
                        if vectors:
                            vectors = np.vstack(vectors)
                            self._build_index(vectors)
                        spinner.text = "向量索引重建完成，加载 {len(self.documents)} 个文档片段"
                        spinner.ok("✅")
                                
            except Exception as e:
                PrettyOutput.print(f"加载缓存索引失败: {str(e)}", 
                                output_type=OutputType.WARNING)
                self.documents = []
                self.index = None
                self.flat_index = None
                self.file_md5_cache = {}

    def _save_cache(self, file_path: str, documents: List[Document], vectors: np.ndarray):
        """Save cache for a single file
        
        Args:
            file_path: File path
            documents: List of documents
            vectors: Document vectors
        """
        try:
            # 保存文件缓存
            cache_path = self._get_cache_path(file_path)
            cache_data = {
                "documents": documents,
                "vectors": vectors
            }
            with lzma.open(cache_path, 'wb') as f:
                pickle.dump(cache_data, f)
                
            # 更新并保存索引
            index_path = os.path.join(self.data_dir, "index.pkl")
            index_data = {
                "file_md5_cache": self.file_md5_cache
            }
            with lzma.open(index_path, 'wb') as f:
                pickle.dump(index_data, f)
                            
        except Exception as e:
            PrettyOutput.print(f"保存缓存失败: {str(e)}", output_type=OutputType.ERROR)

    def _build_index(self, vectors: np.ndarray):
        """Build FAISS index"""
        if vectors.shape[0] == 0:
            self.index = None
            self.flat_index = None
            return
            
        # Create a flat index to store original vectors, for reconstruction
        self.flat_index = faiss.IndexFlatIP(self.vector_dim)
        self.flat_index.add(vectors) # type: ignore
        
        # Create an IVF index for fast search
        nlist = max(4, int(vectors.shape[0] / 1000))  # 每1000个向量一个聚类中心
        quantizer = faiss.IndexFlatIP(self.vector_dim)
        self.index = faiss.IndexIVFFlat(quantizer, self.vector_dim, nlist, faiss.METRIC_INNER_PRODUCT)
        
        # Train and add vectors
        self.index.train(vectors) # type: ignore
        self.index.add(vectors) # type: ignore
        # Set the number of clusters to probe during search
        self.index.nprobe = min(nlist, 10)

    def _split_text(self, text: str) -> List[str]:
        """Use a more intelligent splitting strategy"""
        # Add overlapping blocks to maintain context consistency
        overlap_size = min(200, self.max_paragraph_length // 4)
        
        paragraphs = []
        current_chunk = []
        current_length = 0
        
        # First split by sentence
        sentences = []
        current_sentence = []
        sentence_ends = {'。', '！', '？', '…', '.', '!', '?'}
        
        for char in text:
            current_sentence.append(char)
            if char in sentence_ends:
                sentence = ''.join(current_sentence)
                if sentence.strip():
                    sentences.append(sentence)
                current_sentence = []
        
        if current_sentence:
            sentence = ''.join(current_sentence)
            if sentence.strip():
                sentences.append(sentence)
        
        # Build overlapping blocks based on sentences
        for sentence in sentences:
            if current_length + len(sentence) > self.max_paragraph_length:
                if current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    if len(chunk_text) >= self.min_paragraph_length:
                        paragraphs.append(chunk_text)
                        
                    # Keep some content as overlap
                    overlap_text = ' '.join(current_chunk[-2:])  # Keep the last two sentences
                    current_chunk = []
                    if overlap_text:
                        current_chunk.append(overlap_text)
                        current_length = len(overlap_text)
                    else:
                        current_length = 0
                        
            current_chunk.append(sentence)
            current_length += len(sentence)
        
        # Process the last chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) >= self.min_paragraph_length:
                paragraphs.append(chunk_text)
        
        return paragraphs


    def _process_document_batch(self, documents: List[Document]) -> np.ndarray:
        """Process a batch of documents using shared memory"""
        try:
            texts = []
            self.documents = []  # Reset documents to store chunks
            
            for doc in documents:
                # Split original document into chunks
                chunks = self._split_text(doc.content)
                for chunk_idx, chunk in enumerate(chunks):
                    # Create new Document for each chunk
                    new_metadata = doc.metadata.copy()
                    new_metadata.update({
                        'chunk_index': chunk_idx,
                        'total_chunks': len(chunks),
                        'original_length': len(doc.content)
                    })
                    self.documents.append(Document(
                        content=chunk,
                        metadata=new_metadata,
                        md5=doc.md5
                    ))
                    texts.append(f"File:{doc.metadata['file_path']} Chunk:{chunk_idx} Content:{chunk}")
            
            return get_embedding_batch(self.embedding_model, texts)
        except Exception as e:
            PrettyOutput.print(f"批量处理失败: {str(e)}", OutputType.ERROR)
            return np.zeros((0, self.vector_dim), dtype=np.float32) # type: ignore

    def _process_file(self, file_path: str) -> List[Document]:
        """Process a single file"""
        try:
            # Calculate file MD5
            current_md5 = get_file_md5(file_path)
            if not current_md5:
                return []

            # Check if the file needs to be reprocessed
            if file_path in self.file_md5_cache and self.file_md5_cache[file_path] == current_md5:
                return []

            # Find the appropriate processor
            processor = None
            for p in self.file_processors:
                if p.can_handle(file_path):
                    processor = p
                    break
                    
            if not processor:
                # If no appropriate processor is found, return an empty document
                return []
            
            # Extract text content
            content = processor.extract_text(file_path)
            if not content.strip():
                return []
            
            # Split text
            chunks = self._split_text(content)
            
            # Create document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    content=chunk,
                    metadata={
                        "file_path": file_path,
                        "file_type": Path(file_path).suffix.lower(),
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    },
                    md5=current_md5
                )
                documents.append(doc)
            
            # Update MD5 cache
            self.file_md5_cache[file_path] = current_md5
            return documents
            
        except Exception as e:
            PrettyOutput.print(f"处理文件失败: {file_path}: {str(e)}", 
                            output_type=OutputType.ERROR)
            return []

    def build_index(self, dir: str):
        """Build document index with optimized processing"""
        # Get all files
        with yaspin(text="获取所有文件...", color="cyan") as spinner:
            all_files = []
            for root, _, files in os.walk(dir):
                # Skip .jarvis directories and other ignored paths
                if any(ignored in root for ignored in ['.git', '__pycache__', 'node_modules', '.jarvis']) or \
                any(part.startswith('.jarvis-') for part in root.split(os.sep)):
                    continue
                    
                for file in files:
                    # Skip .jarvis files
                    if '.jarvis' in root:
                        continue
                        
                    file_path = os.path.join(root, file)
                    if os.path.getsize(file_path) > 100 * 1024 * 1024:  # 100MB
                        PrettyOutput.print(f"Skip large file: {file_path}", 
                                        output_type=OutputType.WARNING)
                        continue
                    all_files.append(file_path)
            spinner.text = f"获取所有文件完成，共 {len(all_files)} 个文件"
            spinner.ok("✅")

        # Clean up cache for deleted files
        with yaspin(text="清理缓存...", color="cyan") as spinner:
            deleted_files = set(self.file_md5_cache.keys()) - set(all_files)
            for file_path in deleted_files:
                del self.file_md5_cache[file_path]
                # Remove related documents
                self.documents = [doc for doc in self.documents if doc.metadata['file_path'] != file_path]
            spinner.text = f"清理缓存完成，共 {len(deleted_files)} 个文件"
            spinner.ok("✅")

        # Check file changes
        with yaspin(text="检查文件变化...", color="cyan") as spinner:
            files_to_process = []
            unchanged_files = []
            for file_path in all_files:
                current_md5 = get_file_md5(file_path)
                if current_md5:  # Only process files that can successfully calculate MD5
                    if file_path in self.file_md5_cache and self.file_md5_cache[file_path] == current_md5:
                        # File未变化，记录但不重新处理
                        unchanged_files.append(file_path)
                    else:
                        # New file or modified file
                        files_to_process.append(file_path)
                        spinner.write(f"⚠️ 文件变化: {file_path}")
            spinner.text = f"检查文件变化完成，共 {len(files_to_process)} 个文件需要处理"
            spinner.ok("✅")

        # Keep documents for unchanged files
        unchanged_documents = [doc for doc in self.documents 
                            if doc.metadata['file_path'] in unchanged_files]

        # Process files one by one with optimized vectorization
        if files_to_process:
            new_documents = []
            new_vectors = []
            
            
            for file_path in files_to_process:
                with yaspin(text=f"处理文件 {file_path} ...", color="cyan") as spinner:
                    try:
                        # Process single file
                        file_docs = self._process_file(file_path)
                        if file_docs:
                            # Vectorize documents from this file
                            texts_to_vectorize = [
                                f"File:{doc.metadata['file_path']} Content:{doc.content}"
                                for doc in file_docs
                            ]
                            file_vectors = get_embedding_batch(self.embedding_model, texts_to_vectorize)
                            
                            # Save cache for this file
                            self._save_cache(file_path, file_docs, file_vectors)
                            
                            # Accumulate documents and vectors
                            new_documents.extend(file_docs)
                            new_vectors.append(file_vectors)

                            spinner.text = f"处理文件 {file_path} 完成"
                            spinner.ok("✅")
                            
                    except Exception as e:
                        spinner.text = f"处理文件失败: {file_path}: {str(e)}"
                        spinner.fail("❌")
                    
                    
            # Update documents list
            self.documents.extend(new_documents)

            # Build final index
            if new_vectors:
                with yaspin(text="构建最终索引...", color="cyan") as spinner:
                    all_new_vectors = np.vstack(new_vectors)
                    
                    if self.flat_index is not None:
                        # Get vectors for unchanged documents
                        unchanged_vectors = self._get_unchanged_vectors(unchanged_documents)
                        if unchanged_vectors is not None:
                            final_vectors = np.vstack([unchanged_vectors, all_new_vectors])
                        else:
                            final_vectors = all_new_vectors
                    else:
                        final_vectors = all_new_vectors

                    # Build index
                    spinner.text = f"构建索引..."
                    self._build_index(final_vectors)
                    spinner.text = f"索引构建完成，共 {len(self.documents)} 个文档 "
                    spinner.ok("✅")

            PrettyOutput.print(
                f"索引 {len(self.documents)} 个文档 "
                f"(新/修改: {len(new_documents)}, "
                f"不变: {len(unchanged_documents)})", 
                OutputType.SUCCESS
            )

    def _get_unchanged_vectors(self, unchanged_documents: List[Document]) -> Optional[np.ndarray]:
        """Get vectors for unchanged documents from existing index"""
        try:
            if not unchanged_documents or self.flat_index is None:
                return None

            unchanged_vectors = []
            for doc in unchanged_documents:
                doc_idx = next((i for i, d in enumerate(self.documents) 
                            if d.metadata['file_path'] == doc.metadata['file_path']), None)
                if doc_idx is not None:
                    vector = np.zeros((1, self.vector_dim), dtype=np.float32) # type: ignore
                    self.flat_index.reconstruct(doc_idx, vector.ravel())
                    unchanged_vectors.append(vector)

            return np.vstack(unchanged_vectors) if unchanged_vectors else None
            
        except Exception as e:
            PrettyOutput.print(f"获取不变向量失败: {str(e)}", OutputType.ERROR)
            return None

    def search(self, query: str, top_k: int = 30) -> List[Tuple[Document, float]]:
        """Search documents with context window"""
        if not self.index:
            self.build_index(self.root_dir)
            
        # Get query vector
        with yaspin(text="获取查询向量...", color="cyan") as spinner:
            query_vector = get_embedding(self.embedding_model, query)
            query_vector = query_vector.reshape(1, -1)
            spinner.text = "查询向量获取完成"
            spinner.ok("✅")
        
        # Search with more candidates
        with yaspin(text="搜索...", color="cyan") as spinner:
            initial_k = min(top_k * 4, len(self.documents))
            distances, indices = self.index.search(query_vector, initial_k) # type: ignore
            spinner.text = "搜索完成"
            spinner.ok("✅")
        
        # Process results with context window
        with yaspin(text="处理结果...", color="cyan") as spinner:
            results = []
            seen_files = set()
            
            for idx, dist in zip(indices[0], distances[0]):
                if idx != -1:
                    doc = self.documents[idx]
                    similarity = 1.0 / (1.0 + float(dist))
                    if similarity > 0.3:
                        file_path = doc.metadata['file_path']
                        if file_path not in seen_files:
                            seen_files.add(file_path)
                            
                            # Get full context from original document
                            original_doc = next((d for d in self.documents 
                                            if d.metadata['file_path'] == file_path), None)
                            if original_doc:
                                window_docs = []  # Add this line to initialize the list
                                full_content = original_doc.content
                                # Find all chunks from this file
                                file_chunks = [d for d in self.documents 
                                            if d.metadata['file_path'] == file_path]
                                # Add all related chunks
                                for chunk_doc in file_chunks:
                                    window_docs.append((chunk_doc, similarity * 0.9))
                            
                            results.extend(window_docs)
                            if len(results) >= top_k * (2 * self.context_window + 1):
                                break
            spinner.text = "处理结果完成"
            spinner.ok("✅")
        
        # Sort by similarity and deduplicate
        with yaspin(text="排序...", color="cyan") as spinner:
            results.sort(key=lambda x: x[1], reverse=True)
            seen = set()
            final_results = []
            for doc, score in results:
                key = (doc.metadata['file_path'], doc.metadata['chunk_index'])
                if key not in seen:
                    seen.add(key)
                    final_results.append((doc, score))
                    if len(final_results) >= top_k:
                        break
            spinner.text = "排序完成"
            spinner.ok("✅")
                    
        return final_results

    def query(self, query: str) -> List[Document]:
        """Query related documents
        
        Args:
            query: Query text
            
        Returns:
            List[Document]: Related documents
        """
        results = self.search(query)
        return [doc for doc, _ in results]

    def ask(self, question: str) -> Optional[str]:
        """Ask questions about documents with enhanced context building"""
        try:
            results = self.search(question)
            if not results:
                return None
            
            prompt = f"""
# 🤖 Role Definition
You are a document analysis expert who provides accurate and comprehensive answers based on provided documents.

# 🎯 Core Responsibilities
- Analyze document fragments thoroughly
- Answer questions accurately
- Reference source documents
- Identify missing information
- Maintain professional tone

# 📋 Answer Requirements
## Content Quality
- Base answers strictly on provided documents
- Be specific and precise
- Include relevant quotes when helpful
- Indicate any information gaps
- Use professional language

## Answer Structure
1. Direct Answer
   - Clear and concise response
   - Based on document evidence
   - Professional terminology

2. Supporting Details
   - Relevant document quotes
   - File references
   - Context explanation

3. Information Gaps (if any)
   - Missing information
   - Additional context needed
   - Potential limitations

# 🔍 Analysis Context
Question: {question}

Relevant Documents (by relevance):
"""

            # Add context with length control
            with yaspin(text="添加上下文...", color="cyan") as spinner:
                available_count = self.max_token_count - get_context_token_count(prompt) - 1000
                current_count = 0
                
                for doc, score in results:
                    doc_content = f"""
    ## Document Fragment [Score: {score:.3f}]
    Source: {doc.metadata['file_path']}
    ```
    {doc.content}
    ```
    ---
    """
                    if current_count + get_context_token_count(doc_content) > available_count:
                        PrettyOutput.print(
                            "由于上下文长度限制，部分内容被省略",
                            output_type=OutputType.WARNING
                        )
                        break
                        
                    prompt += doc_content
                    current_count += get_context_token_count(doc_content)

                prompt += """
    # ❗ Important Rules
    1. Only use provided documents
    2. Be precise and accurate
    3. Quote sources when relevant
    4. Indicate missing information
    5. Maintain professional tone
    6. Answer in user's language
    """
                spinner.text = "添加上下文完成"
                spinner.ok("✅")

            with yaspin(text="回答...", color="cyan") as spinner:
                model = PlatformRegistry.get_global_platform_registry().get_normal_platform()
                response = model.chat_until_success(prompt)
                spinner.text = "回答完成"
                spinner.ok("✅")
                return response
            
        except Exception as e:
            PrettyOutput.print(f"回答失败：{str(e)}", OutputType.ERROR)
            return None

    def is_index_built(self) -> bool:
        """Check if the index is built and valid
        
        Returns:
            bool: True if index is built and valid
        """
        return self.index is not None and len(self.documents) > 0

def main():
    """Main function"""
    import argparse
    import sys
    
    # Set standard output encoding to UTF-8
    if sys.stdout.encoding != 'utf-8':
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    
    parser = argparse.ArgumentParser(description='Document retrieval and analysis tool')
    parser.add_argument('--dir', type=str, help='Directory to process')
    parser.add_argument('--build', action='store_true', help='Build document index')
    parser.add_argument('--search', type=str, help='Search document content')
    parser.add_argument('--ask', type=str, help='Ask about documents')
    args = parser.parse_args()

    try:
        current_dir = os.getcwd()
        rag = RAGTool(current_dir)

        if not args.dir:
            args.dir = current_dir

        if args.dir and args.build:
            PrettyOutput.print(f"正在处理目录: {args.dir}", output_type=OutputType.INFO)
            rag.build_index(args.dir)
            return 0

        if args.search or args.ask:

            if args.search:
                results = rag.query(args.search)
                if not results:
                    PrettyOutput.print("未找到相关内容", output_type=OutputType.WARNING)
                    return 1
                    
                for doc in results:
                    output = f"""文件: {doc.metadata['file_path']}\n"""
                    output += f"""片段 {doc.metadata['chunk_index'] + 1}/{doc.metadata['total_chunks']}\n"""
                    output += f"""内容:\n{doc.content}\n"""
                    PrettyOutput.print(output, output_type=OutputType.INFO, lang="markdown")
                return 0

            if args.ask:
                # Call ask method
                response = rag.ask(args.ask)
                if not response:
                    PrettyOutput.print("获取答案失败", output_type=OutputType.WARNING)
                    return 1
                    
                # Display answer
                output = f"""答案:\n{response}"""
                PrettyOutput.print(output, output_type=OutputType.INFO)
                return 0

        PrettyOutput.print("请指定操作参数。使用 -h 查看帮助。", output_type=OutputType.WARNING)
        return 1

    except Exception as e:
        PrettyOutput.print(f"执行失败: {str(e)}", output_type=OutputType.ERROR)
        return 1

if __name__ == "__main__":
    main()
