import io
import threading
import weakref
from typing import Optional, Any, Tuple, Union
from venv import logger

import mammoth
import markdown
from cacherator import JSONCache, Cached
from docx import Document
from docx.document import Document as DocxDocument
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from html2docx import html2docx
from logorator import Logger

from .auth import GoogleAuth
from .exceptions import DocumentNotFoundError, DocumentCreationError, DocumentSaveError
import re

class Docorator(JSONCache):
    def __init__(self, keyfile_path: str, email: Optional[str], document_name: str, clear_cache: bool = True):
        super().__init__(data_id=f"{document_name}", directory="data/docorator", clear_cache=clear_cache)

        self.keyfile_path = keyfile_path
        self.email = email
        self.document_name = document_name

        self.auth = GoogleAuth(keyfile_path)
        self.docs_service = self.auth.get_docs_service()
        self.drive_service = self.auth.get_drive_service()

        if not hasattr(self, "document_id"):
            self.document_id = None

        # Thread-related attributes
        self._load_thread: Optional[threading.Thread] = None
        self._save_thread: Optional[threading.Thread] = None
        self._load_result: Optional[Tuple[Any, Optional[Exception]]] = None
        self._save_result: Optional[Tuple[bool, Optional[Exception]]] = None
        self._thread_lock = threading.Lock()

        self.load()
        weakref.finalize(self, self.close)

    def __str__(self):
        return f"{self.document_name} ({self.url})"

    def __repr__(self):
        return self.__str__()

    @property
    def url(self):
        if self.document_id:
            return f"https://docs.google.com/document/d/{self.document_id}/edit"
        return None

    def _find_document(self) -> bool:
        query = f"name = '{self.document_name}' and mimeType = 'application/vnd.google-apps.document'"
        results = self.drive_service.files().list(q=query, fields="files(id, name)").execute()
        files = results.get('files', [])

        if not files:
            return False

        self.document_id = files[0]['id']

        Logger.note(f"Document found: {self.url}")
        return True

    def _share_with_mail(self):
        if self.email:
            anyone_permission = {
                'type': 'anyone',
                'role': 'writer'
            }
            self.drive_service.permissions().create(
                fileId=self.document_id,
                body=anyone_permission
            ).execute()

            try:
                email_permission = {
                        'type'        : 'user',
                        'role'        : 'writer',
                        'emailAddress': self.email
                }
                self.drive_service.permissions().create(
                        fileId=self.document_id,
                        body=email_permission,
                        sendNotificationEmail=False,
                        fields='id'
                ).execute()
            except Exception as e:
                Logger.note(f"Warning: Failed to share document with {self.email}: {str(e)}")

    def _create_document(self) -> None:
        try:
            file_metadata = {
                'name'    : self.document_name,
                'mimeType': 'application/vnd.google-apps.document'
            }
            document = self.drive_service.files().create(body=file_metadata).execute()
            self.document_id = document['id']

            self._share_with_mail()

            Logger.note("Document created")
        except Exception as e:
            raise DocumentCreationError(f"Failed to create document: {str(e)}")

    def _export_to_docx(self) -> DocxDocument:
        try:
            if not self.document_id:
                raise DocumentNotFoundError("Document ID not found")

            request = self.drive_service.files().export_media(
                fileId=self.document_id,
                mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            file = io.BytesIO()
            downloader = MediaIoBaseDownload(file, request)

            done = False
            while not done:
                status, done = downloader.next_chunk()

            file.seek(0)
            return Document(file)
        except Exception as e:
            raise DocumentNotFoundError(f"Failed to export document to DOCX: {str(e)}")

    def _load_implementation(self) -> DocxDocument:
        if self.document_id:
            try:
                return self._export_to_docx()
            except DocumentNotFoundError:
                self.document_id = None

        if self._find_document():
            return self._export_to_docx()

        self._create_document()
        return Document()

    @Logger()
    def load(self) -> None:
        with self._thread_lock:
            if self._load_thread and self._load_thread.is_alive():
                return

            self._load_result = None

            def thread_target():
                result = None
                exception = None
                try:
                    result = self._load_implementation()
                except Exception as e:
                    exception = e

                with self._thread_lock:
                    self._load_result = (result, exception)

            self._load_thread = threading.Thread(target=thread_target)
            self._load_thread.daemon = True
            self._load_thread.start()

    @Logger()
    def wait_for_load(self) -> DocxDocument:
        if not self._load_thread or not self._load_thread.is_alive() and self._load_result is None:
            self.load()

        if self._load_thread:
            self._load_thread.join()

        if self._load_result:
            result, exception = self._load_result
            if exception:
                raise exception
            return result

        return self._load_implementation()

    @Cached()
    @Logger()
    def as_markdown(self):
        buffer = io.BytesIO()
        doc = self.wait_for_load()
        doc.save(buffer)
        buffer.seek(0)
        result = mammoth.convert_to_markdown(buffer)

        def replace_base64_images_with_alt_tag(text: str) -> str:
            # Capture the alt text inside ![Alt text]... then match a data:image/... pattern
            pattern = r'!\[([^\]]*)\]\(data:image\/[^)]+\)'
            # The replacement uses the captured alt text group (\1) in a bracketed [Image: ...] format
            replacement = r'[Image: \1]'

            return re.sub(pattern, replacement, text)

        markdown_text = replace_base64_images_with_alt_tag(result.value)

        return markdown_text

    def _convert_markdown_to_html(self, md: str = ""):
        html = markdown.markdown(
            md,
            extensions=[
                'markdown.extensions.tables',
                'markdown.extensions.fenced_code',
                'markdown.extensions.codehilite',
                'markdown.extensions.nl2br',
                'markdown.extensions.smarty',
                'markdown.extensions.toc'
            ]
        )
        html = self._replace_images_with_alt_text(html)
        full_html = f"""
                <!DOCTYPE html>
                <html>
                <body>
                    {html}
                </body>
                </html>
                """
        return full_html

    def _replace_images_with_alt_text(self, html: str) -> str:
        def replacer(match):
            alt_text = "[Image]"
            if match.group(1):
                alt_text = f"[Image: {match.group(1)}]"
            return f"<span class='image-placeholder'>{alt_text}</span>"

        # Pattern to match img tags and extract alt attribute
        img_pattern = r'<img.*?alt=["\'](.*?)["\'].*?>'
        html_with_alts = re.sub(img_pattern, replacer, html)

        # Handle imgs without alt attribute
        img_pattern_no_alt = r'<img.*?>'
        return re.sub(img_pattern_no_alt, "<span class='image-placeholder'>[Image]</span>", html_with_alts)

    def _convert_markdown_to_docx(self, md: str = ""):
        docx_bytes = html2docx(self._convert_markdown_to_html(md=md), title=self.document_name)
        return Document(docx_bytes)

    @Logger()
    def _save_implementation(self, document: Union[str, DocxDocument] = "") -> bool:
        if self.document_id is None:
            Logger.note("No Document ID was provided")
            return False
        if isinstance(document, str):
            document = self._convert_markdown_to_docx(document)
        try:
            file_buffer = io.BytesIO()
            document.save(file_buffer)
            file_buffer.seek(0)

            media = MediaIoBaseUpload(
                file_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                resumable=True
            )

            self.drive_service.files().update(
                fileId=self.document_id,
                media_body=media
            ).execute()
            self._share_with_mail()
            return True
        except Exception as e:
            raise DocumentSaveError(f"Failed to save document: {str(e)}")

    @Logger()
    def save(self, document: Union[str, DocxDocument] = "") -> None:
        with self._thread_lock:
            if self._save_thread and self._save_thread.is_alive():
                return

            self._save_result = None

            def thread_target():
                result = False
                exception = None
                try:
                    result = self._save_implementation(document)
                except Exception as e:
                    exception = e

                with self._thread_lock:
                    self._save_result = (result, exception)

            # Start the thread
            self._save_thread = threading.Thread(target=thread_target)
            self._save_thread.daemon = True
            self._save_thread.start()

    @Logger()
    def wait_for_save(self) -> bool:
        if self._save_thread:
            self._save_thread.join()

        if self._save_result:
            result, exception = self._save_result
            if exception:
                raise exception
            return result

        return False

    def close(self):
        self.wait_for_save()
